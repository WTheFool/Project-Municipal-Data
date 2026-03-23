from docx import Document
from docx.shared import Inches
from typing import Dict, Any
import tempfile
import base64
from io import BytesIO

class WordExporter:
    """
    Exports report + graphs to Word (.docx)
    Returns temporary file path
    """

    def export(self, report: str, graphs: Dict[str, str], stats: Dict[str, Any]) -> str:
        doc = Document()
        doc.add_heading("Municipal Data Report", level=0)

        # Add report text
        doc.add_heading("Report Summary", level=1)
        for line in report.split("\n"):
            doc.add_paragraph(line)

        # Add statistics
        doc.add_heading("Statistics Summary", level=1)
        for stat, val in stats.items():
            doc.add_paragraph(f"{stat}: {val}")

        # Add graphs as images (base64)
        doc.add_heading("Graphs", level=1)
        for name, img_b64 in graphs.items():
            img_bytes = BytesIO(base64.b64decode(img_b64))
            doc.add_paragraph(name)
            doc.add_picture(img_bytes, width=Inches(5))

        temp_file = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
        doc.save(temp_file.name)
        return temp_file.name